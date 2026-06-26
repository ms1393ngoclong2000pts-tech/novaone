from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Bao_cao_du_an_NovaOne_Admin.docx")


def set_font(run, name="Calibri", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E1EA", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths):
    set_table_width(table, widths)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            if row_index == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_font(run, size=10, bold=True, color="1F4D78")
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font(run, size=10)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=16 if level == 1 else 10, after=6)
    run = paragraph.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color="2E74B5")
    elif level == 2:
        set_font(run, size=13, bold=True, color="2E74B5")
    else:
        set_font(run, size=12, bold=True, color="1F4D78")
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    set_paragraph_spacing(paragraph, after=6, line=1.1)
    for run in paragraph.runs:
        set_font(run, size=11)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, after=4, line=1.167)
        run = paragraph.add_run(item)
        set_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(paragraph, after=4, line=1.167)
        run = paragraph.add_run(item)
        set_font(run, size=11)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(title + ": ")
    set_font(run, size=10.5, bold=True, color="1F4D78")
    run = paragraph.add_run(text)
    set_font(run, size=10.5)
    doc.add_paragraph()


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Hạng mục"
    table.rows[0].cells[1].text = "Nội dung"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    style_table(table, [1.85, 4.65])
    for row in table.rows[1:]:
        for run in row.cells[0].paragraphs[0].runs:
            set_font(run, size=10, bold=True, color="1F4D78")
    return table


def add_feature_table(doc):
    rows = [
        ("Đăng nhập, profile, đổi mật khẩu", "AuthController, PHP session, CSRF token, lưu hồ sơ trong storage/data.json."),
        ("Dashboard ứng dụng", "CSS grid, route dashboard, hiệu ứng hover bằng transition/transform."),
        ("Sidebar động", "Layout app.php, JavaScript toggle class, CSS responsive drawer."),
        ("Thông báo hoạt động", "Helper add_notification(), NotificationController, dữ liệu _notifications trong JSON."),
        ("Nhân viên", "EmployeeController, import Excel qua XlsxReader, export CSV."),
        ("Hợp đồng lao động", "ContractController, tab loại hợp đồng, validate ngày, modal CRUD."),
        ("Bảng lương", "PayrollController, lọc bộ phận/ngày, hoàn thành bảng lương, export CSV."),
        ("Bảo hiểm xã hội", "SocialInsuranceController, chọn nhân viên từ employees, tự tính đóng BHXH."),
        ("Phiếu yêu cầu", "RequestFormController, loại phiếu, trạng thái phê duyệt, modal chi tiết."),
        ("Vi phạm và khen thưởng", "ViolationController, RewardController, liên kết người nhận với danh sách nhân viên."),
        ("Dự án, công việc, báo cáo hằng ngày", "Module Work: ProjectController, WorkItemController, DailyReportController."),
        ("Nhà cung cấp", "Module Business: SupplierController, import Excel/CSV, CRUD và modal thông tin."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Tính năng"
    table.rows[0].cells[1].text = "Điểm đặc biệt"
    table.rows[0].cells[2].text = "Công nghệ / cách làm"
    for feature, tech in rows:
        cells = table.add_row().cells
        cells[0].text = feature
        cells[1].text = "Có giao diện riêng, dữ liệu mẫu, tìm kiếm, phân trang, thông báo thao tác."
        cells[2].text = tech
    style_table(table, [1.8, 2.2, 2.5])
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=0, after=6)
    run = title.add_run("BÁO CÁO DỰ ÁN NOVAONE ADMIN")
    set_font(run, size=20, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=14)
    run = subtitle.add_run("Ứng dụng web quản trị doanh nghiệp bằng PHP thuần")
    set_font(run, size=12, color="555555")

    add_callout(
        doc,
        "Tóm tắt",
        "NovaOne Admin là hệ thống quản trị nội bộ dạng ERP mini, tập trung vào nhân sự, công việc, kinh doanh, kho, báo cáo và hệ thống. Ứng dụng đã được tổ chức theo module, có giao diện responsive và các màn nghiệp vụ có chung trải nghiệm lọc, tìm kiếm, phân trang, modal chi tiết và thông báo hoạt động.",
    )

    add_heading(doc, "1. Tổng quan dự án")
    add_body(
        doc,
        "Dự án NovaOne Admin được xây dựng để mô phỏng một hệ thống quản trị doanh nghiệp có thể vận hành nhiều nghiệp vụ trong cùng một giao diện. Ứng dụng hiện chạy local bằng PHP, sử dụng dữ liệu JSON để thuận tiện demo, kiểm thử và mở rộng chức năng nhanh.",
    )
    add_key_value_table(
        doc,
        [
            ("Loại hệ thống", "Ứng dụng web quản trị doanh nghiệp"),
            ("Ngôn ngữ chính", "PHP thuần, HTML, CSS, JavaScript"),
            ("Kiến trúc", "MVC đơn giản, chia module theo nghiệp vụ"),
            ("Lưu trữ dữ liệu", "File JSON tại storage/data.json"),
            ("Mục tiêu", "Demo đầy đủ luồng quản trị, dễ mở rộng sang database thật"),
        ],
    )

    add_heading(doc, "2. Công nghệ và thành phần sử dụng")
    add_bullets(
        doc,
        [
            "PHP thuần: xử lý route, controller, session, CSRF, import/export và lưu dữ liệu.",
            "HTML/CSS/JavaScript: xây dựng layout, sidebar động, modal, responsive và các tương tác giao diện.",
            "JSON DataStore: lưu dữ liệu demo trong storage/data.json, dùng class DataStore để đọc/ghi tập trung.",
            "XlsxReader tự viết: đọc file .xlsx không cần thư viện ngoài, dùng cho import nhân viên, dự án và nhà cung cấp.",
            "CSV export/import: xuất dữ liệu ra file CSV và hỗ trợ import CSV cho nhà cung cấp.",
            "Module architecture: HumanResources, Work, Business giúp source rõ ràng, dễ tìm và dễ bảo trì.",
        ],
    )

    add_heading(doc, "3. Cấu trúc source")
    add_key_value_table(
        doc,
        [
            ("public/index.php", "Điểm vào ứng dụng, định nghĩa các route chính."),
            ("app/Core", "Helper, View renderer, DataStore, XlsxReader."),
            ("app/Controllers", "Controller dùng chung: Auth, Dashboard, Search, Notification, Resource."),
            ("app/Modules/HumanResources", "Nhân viên, hợp đồng, bảng lương, BHXH, phiếu yêu cầu, vi phạm, khen thưởng."),
            ("app/Modules/Work", "Dự án, danh sách công việc, báo cáo hằng ngày."),
            ("app/Modules/Business", "Nhà cung cấp và các chức năng kinh doanh mở rộng."),
            ("public/assets/app.css", "Toàn bộ giao diện, responsive, bảng, modal, dashboard và sidebar."),
            ("storage/data.json", "Dữ liệu runtime của bản demo."),
        ],
    )

    add_heading(doc, "4. Các tính năng đã triển khai")
    add_feature_table(doc)

    add_heading(doc, "5. Các tính năng đặc biệt")
    add_bullets(
        doc,
        [
            "Không dùng framework nhưng source vẫn được chia module tương đối chuẩn, dễ mở rộng và dễ kiểm soát.",
            "Import Excel được xử lý bằng class XlsxReader tự viết, giảm phụ thuộc thư viện bên ngoài.",
            "Thông báo hoạt động được tạo tự động sau các thao tác thêm, sửa, xóa, import hoặc thay đổi trạng thái.",
            "Các màn nghiệp vụ có UX đồng nhất: header, filter, search, table, pagination, modal chi tiết và flash message.",
            "Nhiều màn liên kết với danh sách nhân viên, ví dụ BHXH, vi phạm, khen thưởng, phiếu yêu cầu và báo cáo công việc.",
            "Responsive được tối ưu cho mobile/tablet: sidebar dạng drawer, bảng chuyển thành card, form co về một hoặc hai cột.",
            "Dữ liệu mẫu tự sinh theo từng module để người dùng mở trang là có thể thao tác ngay.",
        ],
    )

    add_heading(doc, "6. Cách các tính năng được thực hiện")
    add_numbered(
        doc,
        [
            "Route trong public/index.php nhận tham số route và gọi controller tương ứng.",
            "Controller đọc dữ liệu qua DataStore, lọc/tìm kiếm/sắp xếp/phân trang rồi truyền sang View.",
            "View hiển thị HTML, form, modal và nhúng JavaScript nhỏ để mở/sửa dữ liệu.",
            "Form POST sử dụng CSRF token, controller validate dữ liệu trước khi ghi vào JSON.",
            "Sau thao tác thành công, helper add_notification() ghi thông báo hoạt động vào storage/data.json.",
            "CSS tại public/assets/app.css định nghĩa layout desktop và các media query cho tablet/mobile.",
        ],
    )

    add_heading(doc, "7. Đánh giá hiện trạng")
    add_body(
        doc,
        "Ứng dụng hiện phù hợp cho demo nghiệp vụ, trình bày giao diện và kiểm thử luồng xử lý nội bộ. Các module chính đã có dữ liệu mẫu và thao tác CRUD. Kiến trúc có thể chuyển từ JSON sang database như MySQL mà vẫn giữ phần lớn controller và view hiện tại.",
    )

    add_heading(doc, "8. Đề xuất phát triển tiếp")
    add_bullets(
        doc,
        [
            "Chuyển DataStore JSON sang database MySQL hoặc PostgreSQL.",
            "Bổ sung phân quyền chi tiết theo vai trò và module.",
            "Tách CSS thành nhiều file theo module hoặc dùng build pipeline khi dự án lớn hơn.",
            "Thêm kiểm thử tự động cho controller, import/export và validate dữ liệu.",
            "Hoàn thiện các phân hệ còn lại: dịch vụ, bán hàng, kho, báo cáo nâng cao và quản lý hệ thống.",
            "Tối ưu bảo mật upload file, giới hạn dung lượng và kiểm tra định dạng chặt chẽ hơn.",
        ],
    )

    add_heading(doc, "9. Kết luận")
    add_body(
        doc,
        "NovaOne Admin đã hình thành nền tảng quản trị doanh nghiệp với giao diện rõ ràng, nhiều phân hệ nghiệp vụ và cách tổ chức source dễ mở rộng. Điểm mạnh của dự án là sử dụng PHP thuần nhưng vẫn giữ được cấu trúc module, có import Excel, thông báo hoạt động, responsive và trải nghiệm quản trị thống nhất.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("NovaOne Admin - Báo cáo dự án")
    set_font(run, size=9, color="666666")

    doc.save(OUT)


if __name__ == "__main__":
    build()
